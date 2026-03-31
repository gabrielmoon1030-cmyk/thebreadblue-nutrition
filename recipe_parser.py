"""
배합비 파서 (Excel + Word)
- 배합비 폴더에서 원재료명 + 투입량(g) 자동 추출
- 다양한 엑셀 형식 대응 (레시피/원가 파일)
- Word(.docx) 배합비율표 지원
"""
import re
import openpyxl
from pathlib import Path
from ingredient_db import normalize_name

# 로컬 전용 (클라우드에서는 파일 업로드 모드만 사용)
_local_recipe_dir = r"C:\Users\Moon\OneDrive\03. 양지희_품질\7. 배합비"
RECIPE_DIR = Path(_local_recipe_dir) if Path(_local_recipe_dir).exists() else None


def parse_recipe_file(filepath):
    """단일 배합비 Excel에서 (제품명, [(원재료, 투입량g), ...]) 추출"""
    results = []
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    except Exception as e:
        return []

    for ws in wb.worksheets:
        sheet_name = ws.title
        # '품목' 시트는 라벨용 재정렬이므로 스킵
        if '품목' in sheet_name:
            continue

        # 제품명 추출 (보통 B1 또는 B2)
        product_name = None
        for row_num in range(1, 4):
            cell_b = ws.cell(row=row_num, column=2).value
            if cell_b and isinstance(cell_b, str) and len(cell_b.strip()) > 1:
                val = cell_b.strip()
                if not any(k in val for k in ['단위', '규격', '원재료', '배합']):
                    product_name = val
                    break

        if not product_name:
            product_name = Path(filepath).stem

        # 원재료 + 투입량 추출
        ingredients = []
        header_row = None

        # 헤더 행 찾기 (원재료 컬럼)
        for row_num in range(1, 10):
            cell_a = ws.cell(row=row_num, column=1).value
            if cell_a and '원재료' in str(cell_a):
                header_row = row_num
                break

        if header_row is None:
            continue

        # 투입량 컬럼 찾기 (1배합 또는 다음 숫자 컬럼)
        qty_col = 2  # 기본값: B열
        for col in range(2, 8):
            cell = ws.cell(row=header_row, column=col).value
            if cell and ('1배합' in str(cell) or '배합' in str(cell)):
                qty_col = col
                break

        # 데이터 읽기
        for row_num in range(header_row + 1, 60):
            cell_a = ws.cell(row=row_num, column=1).value
            cell_qty = ws.cell(row=row_num, column=qty_col).value

            if cell_a is None:
                continue

            name = str(cell_a).strip()

            # 종료 조건
            if any(k in name for k in ['합계', 'MEMO', 'memo', '소계', '총합', '개당', '판매']):
                break

            # 필터
            if len(name) < 2:
                continue
            if re.match(r'^[\d.,\s%()]+$', name):
                continue

            # 공정 지시어 제외 (실제 원재료명은 통과시킴)
            skip_keywords = ['반죽', '성형', '굽기', '냉각', '분할', '℃', '오븐',
                           '뚜껑', '모양', '올려', '넣어', '섞어', '비고',
                           '단위', '규격', '입고', '원가', '백분율', '제품명']
            # '글루텐 60%' 같은 공정 메모는 제외하되, '활성밀글루텐' 같은 원재료는 통과
            if any(k in name for k in skip_keywords):
                continue
            if re.search(r'글루텐\s*\d', name) or name == '글루텐':
                continue

            # 투입량 파싱
            qty = 0
            if cell_qty is not None:
                try:
                    qty = float(cell_qty)
                except (ValueError, TypeError):
                    # 문자열에서 숫자 추출 시도
                    match = re.search(r'[\d.]+', str(cell_qty))
                    if match:
                        qty = float(match.group())

            if qty > 0:
                norm_name = normalize_name(name)
                ingredients.append((norm_name, qty))

        if ingredients:
            results.append({
                "product_name": product_name,
                "sheet_name": sheet_name,
                "file": str(filepath),
                "ingredients": ingredients,
                "total_weight": sum(q for _, q in ingredients)
            })

    wb.close()
    return results


def parse_all_recipes():
    """전체 배합비 폴더 파싱 (로컬 전용)"""
    if RECIPE_DIR is None or not RECIPE_DIR.exists():
        return []
    all_recipes = []
    xlsx_files = list(RECIPE_DIR.rglob("*.xlsx"))

    for fpath in xlsx_files:
        if "~$" in fpath.name:
            continue
        recipes = parse_recipe_file(fpath)
        all_recipes.extend(recipes)

    # 중복 제거 (같은 제품명이면 가장 최근 파일 우선)
    seen = {}
    for r in all_recipes:
        key = r["product_name"]
        if key not in seen:
            seen[key] = r
        # 나중에 파싱된 것이 더 최근 → 덮어쓰기

    return list(seen.values())


def parse_docx_file(filepath, total_weight_g=None):
    """Word(.docx) 배합비율표에서 (제품명, [(원재료, 투입량g), ...]) 추출

    docx는 보통 배합비율(%)로 되어 있으므로 total_weight_g가 필요.
    파일명에서 중량을 자동 추출 시도 (예: '치즈케이크 520g').
    """
    from docx import Document

    if isinstance(filepath, (str, Path)):
        doc = Document(filepath)
        filename = Path(filepath).stem
    else:
        # BytesIO (Streamlit 업로드)
        doc = Document(filepath)
        filename = getattr(filepath, 'name', 'unknown').replace('.docx', '')

    # 파일명에서 총 중량 추출 (예: "치즈케이크 520g")
    if total_weight_g is None:
        wt_match = re.search(r'(\d+)\s*[gG]', filename)
        if wt_match:
            total_weight_g = float(wt_match.group(1))

    # 제품명 추출 (파일명에서)
    product_name = re.sub(r'\d+\.\s*', '', filename)  # "2. 원료성분 및 배합비율(치즈케이크 520g)" → 정리
    name_match = re.search(r'[((](.+?)\s*\d*\s*[gG]?\s*[))]', filename)
    if name_match:
        product_name = name_match.group(1).strip()
    else:
        product_name = re.sub(r'원료성분\s*(및|&)?\s*배합비율\s*', '', product_name).strip()
        if not product_name:
            product_name = filename

    results = []

    for table in doc.tables:
        ingredients = []
        name_col = None
        ratio_col = None
        is_percentage = False

        # 헤더 행에서 컬럼 위치 찾기
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        for idx, header in enumerate(header_cells):
            if '원재료' in header or '원료' in header:
                name_col = idx
            if '배합비율' in header or '비율' in header or '%' in header:
                ratio_col = idx
                is_percentage = True
            if '투입량' in header or '중량' in header or '배합량' in header:
                ratio_col = idx
                is_percentage = False

        if name_col is None or ratio_col is None:
            continue

        # 데이터 행 읽기
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            name = cells[name_col] if name_col < len(cells) else ''
            value_str = cells[ratio_col] if ratio_col < len(cells) else ''

            if not name or len(name) < 2:
                continue
            if any(k in name for k in ['합계', '합 계', '소계', '총합', '계']):
                continue

            # 숫자 파싱
            value_str = value_str.replace(',', '')
            match = re.search(r'[\d.]+', value_str)
            if not match:
                continue
            value = float(match.group())
            if value <= 0:
                continue

            # %이면 g로 변환
            if is_percentage and total_weight_g:
                qty_g = value / 100 * total_weight_g
            elif is_percentage:
                # 총 중량 모르면 100g 기준으로 계산
                qty_g = value  # 비율 그대로 (100g 기준)
            else:
                qty_g = value

            norm_name = normalize_name(name)
            ingredients.append((norm_name, round(qty_g, 2)))

        if ingredients:
            actual_total = total_weight_g if total_weight_g else sum(q for _, q in ingredients)
            results.append({
                "product_name": product_name,
                "sheet_name": "본문",
                "file": str(filepath) if isinstance(filepath, (str, Path)) else filename,
                "ingredients": ingredients,
                "total_weight": actual_total,
                "is_percentage": is_percentage,
                "source_weight_g": total_weight_g,
            })

    return results


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8')

    recipes = parse_all_recipes()
    print(f"총 {len(recipes)}개 레시피 파싱 완료\n")

    for r in recipes[:5]:
        print(f"[{r['product_name']}] (총 {r['total_weight']:.0f}g)")
        for name, qty in r['ingredients'][:8]:
            print(f"  - {name}: {qty:.1f}g")
        if len(r['ingredients']) > 8:
            print(f"  ... 외 {len(r['ingredients'])-8}건")
        print()
